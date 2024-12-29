from django.http import JsonResponse, HttpResponse
from django.views.decorators.csrf import csrf_exempt
from django.utils.encoding import escape_uri_path
from .forms import CreateNoteForm
from .forms import UpdateNoteForm
from .forms import NoteQueryForm
from .forms import FolderQueryForm
from .models import Note, Tag
from user.utils import verify_and_refresh_token  # 自定义的 token 验证函数
from rest_framework_simplejwt.tokens import AccessToken
import json
from django.db.models import Q  # 引入 Q 对象，用于组合多个查询条件
import io
import os
from docx import Document

@csrf_exempt
def create_note_view(request):
    if request.method == 'POST':
        # 验证并刷新 token
        try:
            token = verify_and_refresh_token(request)
            access_token = AccessToken(token)
            user_id = access_token['user_id']  # 从 token 中获取 user_id
        except Exception as e:
            return JsonResponse({"code": 1, "msg": f"Token 验证失败: {str(e)}"})

        # 获取前端请求数据
        title = request.POST.get('title')
        content = request.POST.get('content')
        folder_id = request.POST.get('folder_id')
        tag_names = request.POST.getlist('tags')  # 获取标签名称列表

        if not title:
            title = '新建笔记'

        if not folder_id:
            return JsonResponse({"code": 1, "msg": "所属文件夹不能为空"})

        # 将标签名称转换为标签 ID 列表
        tag_ids = []
        for tag_name in tag_names:
            try:
                tag = Tag.objects.get(name=tag_name)  # 查找标签名称对应的标签实例
                tag_ids.append(tag.id)  # 将标签 ID 加入列表
            except Tag.DoesNotExist:
                tag = Tag.objects.create(name=tag_name)
                tag_ids.append(tag.id)

        # 构建包含所有字段数据的字典
        note_data = {
            'title': title,
            'content': content,
            'folder_id': folder_id,
            'tags': tag_ids  # 传递标签 ID 列表
        }

        # 创建 Note 表单实例并传入数据
        form = CreateNoteForm(note_data)

        if form.is_valid():
            # 保存笔记时，先不要提交到数据库
            note = form.save(commit=False)
            note.user_id = user_id  # 将当前用户关联到笔记
            note.title = title
            note.save()  # 保存笔记实例

            # 关联标签到笔记
            note.tags.set(tag_ids)  # 使用 set() 方法将标签关联到笔记

            # 返回成功的响应
            return JsonResponse({
                "code": 0,
                "msg": "success",
                "data": {
                    "id": note.id,
                    "user_id": note.user_id,
                    "title": note.title,
                    "content": note.content,
                    "folder_id": note.folder_id,
                    "tags": [tag.name for tag in note.tags.all()],  # 返回标签的名称
                    "created_at": note.created_at.strftime("%Y-%m-%d %H:%M:%S"),
                    "updated_at": note.updated_at.strftime("%Y-%m-%d %H:%M:%S"),
                    "deleted_at": note.deleted_at.strftime("%Y-%m-%d %H:%M:%S") if note.deleted_at else None
                }
            })
        else:
            return JsonResponse({"code": 1, "msg": f"创建笔记失败: {form.errors}"})

    return JsonResponse({"code": 1, "msg": "无效的请求方法"})


@csrf_exempt  # 禁用 CSRF 验证
def update_note_view(request):
    if request.method == 'POST':
        # 验证并刷新 token
        try:
            token = verify_and_refresh_token(request)
            access_token = AccessToken(token)
            user_id = access_token['user_id']  # 从 token 中获取 user_id
        except Exception as e:
            return JsonResponse({"code": 1, "msg": f"Token 验证失败: {str(e)}"})

        # 获取前端提交的表单数据
        form = UpdateNoteForm(request.POST)
        if not form.is_valid():
            return JsonResponse({"code": 1, "msg": f"参数验证失败: {form.errors}"})

        # 获取验证后的数据
        note_id = form.cleaned_data['note_id']
        content = form.cleaned_data['content']

        # 获取指定 ID 的笔记
        try:
            note = Note.objects.get(id=note_id)
        except Note.DoesNotExist:
            return JsonResponse({"code": 1, "msg": "笔记不存在"})

        # 检查是否为当前用户的笔记
        if note.user_id != user_id:
            return JsonResponse({"code": 1, "msg": "您没有权限修改此笔记"})

        # 更新笔记内容
        note.content = content
        note.save()

        # 返回更新后的笔记信息
        return JsonResponse({
            "code": 0,
            "msg": "笔记更新成功",
            "data": {
                "id": note.id,
                "user_id": note.user_id,
                "title": note.title,
                "content": note.content,
                "folder_id": note.folder_id,
                "tags": [tag.name for tag in note.tags.all()],
                "created_at": note.created_at.strftime("%Y-%m-%d %H:%M:%S"),
                "updated_at": note.updated_at.strftime("%Y-%m-%d %H:%M:%S"),
                "deleted_at": note.deleted_at.strftime("%Y-%m-%d %H:%M:%S") if note.deleted_at else None
            }
        })

    return JsonResponse({"code": 1, "msg": "无效的请求方法"})

def get_note_view(request):
    if request.method == 'GET':
        # 验证并刷新 token
        try:
            token = verify_and_refresh_token(request)
            access_token = AccessToken(token)
            user_id = access_token['user_id']  # 从 token 中获取 user_id
        except Exception as e:
            return JsonResponse({"code": 1, "msg": f"Token 验证失败: {str(e)}"})

        # 获取 query 参数，并进行表单验证
        form = NoteQueryForm(request.GET)
        if not form.is_valid():
            return JsonResponse({"code": 1, "msg": f"参数验证失败: {form.errors}"})

        # 获取验证后的数据
        note_id = form.cleaned_data['note_id']

        # 获取指定 ID 的笔记
        try:
            note = Note.objects.get(id=note_id)
        except Note.DoesNotExist:
            return JsonResponse({"code": 1, "msg": "笔记不存在"})

        # 检查是否为当前用户的笔记
        if note.user_id != user_id:
            return JsonResponse({"code": 1, "msg": "您没有权限查看此笔记"})

        # 返回笔记内容
        return JsonResponse({
            "code": 0,
            "msg": "笔记获取成功",
            "data": {
                "id": note.id,
                "user_id": note.user_id,
                "title": note.title,
                "content": note.content,
                "folder_id": note.folder_id,
                "tags": [tag.name for tag in note.tags.all()],
                "created_at": note.created_at.strftime("%Y-%m-%d %H:%M:%S"),
                "updated_at": note.updated_at.strftime("%Y-%m-%d %H:%M:%S"),
                "deleted_at": note.deleted_at.strftime("%Y-%m-%d %H:%M:%S") if note.deleted_at else None
            }
        })

    return JsonResponse({"code": 1, "msg": "无效的请求方法"})

@csrf_exempt
def get_notes_in_folder_view(request):
    if request.method == 'GET':
        # 验证并刷新 token
        try:
            token = verify_and_refresh_token(request)
            access_token = AccessToken(token)
            user_id = access_token['user_id']  # 从 token 中获取 user_id
        except Exception as e:
            return JsonResponse({"code": 1, "msg": f"Token 验证失败: {str(e)}"})

        # 获取 query 参数并进行表单验证
        form = FolderQueryForm(request.GET)
        if not form.is_valid():
            return JsonResponse({"code": 1, "msg": f"参数验证失败: {form.errors}"})

        # 获取验证后的数据
        folder_id = form.cleaned_data['folder_id']

        # 获取当前文件夹下所有的笔记
        notes = Note.objects.filter(user_id=user_id, folder_id=folder_id)

        # 构建响应数据
        notes_data = [
            {
                "id": note.id,
                "user_id": note.user_id,
                "title": note.title,
                "folder_id": note.folder_id,
                "created_at": note.created_at.strftime("%Y-%m-%d %H:%M:%S"),
                "tag": [tag.name for tag in note.tags.all()],
            }
            for note in notes
        ]

        # 返回笔记列表
        return JsonResponse({
            "code": 0,
            "msg": "success",
            "data": notes_data
        })

    return JsonResponse({"code": 1, "msg": "无效的请求方法"})

@csrf_exempt
def delete_note_view(request):
    if request.method == 'POST':
        # 验证并刷新 token
        try:
            token = verify_and_refresh_token(request)
            access_token = AccessToken(token)
            user_id = access_token['user_id']  # 从 token 中获取 user_id
        except Exception as e:
            return JsonResponse({"code": 1, "msg": f"Token 验证失败: {str(e)}"})

        # 获取前端提交的 note_id
        note_id = request.POST.get('note_id')
        if not note_id:
            return JsonResponse({"code": 1, "msg": "笔记 ID 为必填项"})

        # 获取指定 ID 的笔记
        try:
            note = Note.objects.get(id=note_id, user_id=user_id)
        except Note.DoesNotExist:
            return JsonResponse({"code": 1, "msg": "笔记不存在或无权限删除"})

        # 删除笔记
        note.delete()

        # 返回删除成功的响应
        return JsonResponse({"code": 0, "msg": "笔记删除成功"})

    return JsonResponse({"code": 1, "msg": "无效的请求方法"})

@csrf_exempt 
def update_note_tag_view(request):
    if request.method == 'POST':  # 使用 POST 方法
        # 验证并刷新 token
        try:
            token = verify_and_refresh_token(request)
            access_token = AccessToken(token)
            user_id = access_token['user_id']  # 从 token 中获取 user_id
        except Exception as e:
            return JsonResponse({"code": 1, "msg": f"Token 验证失败: {str(e)}"})

        # 获取前端提交的数据
        note_id = request.POST.get('note_id')  # 获取笔记 ID
        tag_names = request.POST.getlist('tag')  # 获取标签名称列表（多个标签）

        # # 验证必要参数
        # if not note_id or not tag_names:
        #     return JsonResponse({"code": 1, "msg": "参数缺失：需要 note_id 和 tag"})
                # 验证必要参数
        if not note_id:
            return JsonResponse({"code": 1, "msg": "参数缺失：需要 note_id"})

        # 查询指定的笔记
        try:
            note = Note.objects.get(id=note_id, user_id=user_id)
        except Note.DoesNotExist:
            return JsonResponse({"code": 1, "msg": "笔记不存在或无权限修改"})

        # 清空现有标签
        note.tags.clear()

        # 查询或创建标签并添加到笔记
        for tag_name in tag_names:
            tag, created = Tag.objects.get_or_create(name=tag_name.strip())  # 清除多余空格
            note.tags.add(tag)

        # 构建响应数据
        note_data = {
            "id": note.id,
            "user_id": note.user_id,
            "title": note.title,
            "content": note.content,
            "tags": [tag.name for tag in note.tags.all()],  # 返回标签列表
            "folder_id": note.folder_id,
            "created_at": note.created_at.strftime("%Y-%m-%d %H:%M:%S"),
            "updated_at": note.updated_at.strftime("%Y-%m-%d %H:%M:%S") if note.updated_at else None,
            "deleted_at": note.deleted_at.strftime("%Y-%m-%d %H:%M:%S") if note.deleted_at else None,
        }

        # 返回成功响应
        return JsonResponse({"code": 0, "msg": "标签更新成功", "data": note_data})

    # 如果请求方法不是 POST，则返回错误信息
    return JsonResponse({"code": 1, "msg": "无效的请求方法"})


@csrf_exempt
def search_notes_by_tag_view(request):
    if request.method == 'GET':
        try:
            # 验证并刷新 token
            token = verify_and_refresh_token(request)
            access_token = AccessToken(token)
            user_id = access_token['user_id']
        except Exception as e:
            return JsonResponse({"code": 1, "msg": f"Token 验证失败: {str(e)}"})

        # 获取传递的参数，tag_name 用于传递标签或标题
        search_param = request.GET.get('tag')
        if not search_param:
            return JsonResponse({"code": 1, "msg": "参数缺失：需要 tag 或 title"})

        # 处理分页参数
        def get_int_param(param_value, default_value):
            try:
                return int(param_value) if param_value else default_value
            except ValueError:
                return default_value

        limit = get_int_param(request.GET.get('limit'), 10)
        offset = get_int_param(request.GET.get('offset'), 0)

        # 构建查询条件
        query = Q(user_id=user_id)  # 只查询当前用户的笔记

        # 同时在 tags 和 title 中进行模糊搜索
        query &= (Q(title__icontains=search_param) | Q(tags__name__icontains=search_param))

        # 获取笔记数据，使用 distinct 去重
        notes = Note.objects.filter(query).distinct().order_by('-created_at')[offset:offset + limit]

        # 将笔记数据转换为列表
        notes_data = [
            {
                "id": note.id,
                "user_id": note.user_id,
                "title": note.title,
                "content": note.content,
                "folder_id": note.folder_id,
                "tag": search_param,  # 返回传入的 tag 或 title
                "created_at": note.created_at.strftime("%Y-%m-%d %H:%M:%S"),
                "updated_at": note.updated_at.strftime("%Y-%m-%d %H:%M:%S") if note.updated_at else None,
                "deleted_at": note.deleted_at.strftime("%Y-%m-%d %H:%M:%S") if note.deleted_at else None,
            }
            for note in notes
        ]

        return JsonResponse({
            "code": 0,
            "msg": "success",
            "data": {
                "notes": notes_data,
            }
        })

    return JsonResponse({"code": 1, "msg": "无效的请求方法"})

@csrf_exempt  # 禁用 CSRF 验证
def update_note_title_view(request):
    if request.method == 'POST':
        # 验证并刷新 token
        try:
            token = verify_and_refresh_token(request)
            access_token = AccessToken(token)
            user_id = access_token['user_id']  # 从 token 中获取 user_id
        except Exception as e:
            return JsonResponse({"code": 1, "msg": f"Token 验证失败: {str(e)}"})

        # 获取前端提交的数据
        note_id = request.POST.get('note_id')  # 获取笔记 ID
        title = request.POST.get('title')  # 获取笔记标题

        # 验证必要参数
        if not note_id or not title:
            return JsonResponse({"code": 1, "msg": "参数缺失：需要 note_id 和 title"})

        # 查询指定 ID 的笔记
        try:
            note = Note.objects.get(id=note_id)
        except Note.DoesNotExist:
            return JsonResponse({"code": 1, "msg": "笔记不存在"})

        # 检查是否为当前用户的笔记
        if note.user_id != user_id:
            return JsonResponse({"code": 1, "msg": "您没有权限修改此笔记"})

        # 更新笔记标题
        note.title = title
        note.save()

        # 返回更新后的笔记信息
        return JsonResponse({
            "code": 0,
            "msg": "success",
            "data": {
                "id": note.id,
                "user_id": note.user_id,
                "title": note.title,
                "content": note.content,
                "folder_id": note.folder_id,
                "tags": [tag.name for tag in note.tags.all()],
                "created_at": note.created_at.strftime("%Y-%m-%d %H:%M:%S"),
                "updated_at": note.updated_at.strftime("%Y-%m-%d %H:%M:%S") if note.updated_at else None,
                "deleted_at": note.deleted_at.strftime("%Y-%m-%d %H:%M:%S") if note.deleted_at else None,
            }
        })

    return JsonResponse({"code": 1, "msg": "无效的请求方法"})

# @csrf_exempt
# def export_note_view(request):
#     """
#     导出笔记内容为文件的视图，可选择txt或docx格式
#     请求方法: GET
#     请求参数:
#         note_id (必填)  : 笔记ID
#         format (选填)   : 导出格式, 可选 'txt' 或 'docx'，默认为 'txt'
#     """
#     print("request.method =", request.method)
#     if request.method == 'GET':
#         # 验证并刷新 token
#         try:
#             token = verify_and_refresh_token(request)
#             access_token = AccessToken(token)
#             user_id = access_token['user_id']  # 从 token 中获取 user_id
#         except Exception as e:
#             return JsonResponse({"code": 1, "msg": f"Token 验证失败: {str(e)}"})

#         # 获取前端传递的 note_id 和导出格式
#         note_id = request.GET.get('note_id')
#         export_format = request.GET.get('format', 'txt').lower()  # 默认 txt

#         if not note_id:
#             return JsonResponse({"code": 1, "msg": "缺少参数：note_id"})

#         # 查询指定笔记
#         try:
#             note = Note.objects.get(id=note_id, user_id=user_id)
#         except Note.DoesNotExist:
#             return JsonResponse({"code": 1, "msg": "笔记不存在或无权限导出"})

#         # 准备文件名
#         file_name_base = note.title.strip() if note.title.strip() else "未命名"
#         # 为了避免不合法的文件名，可以再做一次简单处理
#         file_name_base = file_name_base.replace('/', '_').replace('\\', '_').replace(':', '-')

#         # 根据导出格式进行处理
#         if export_format == 'docx':
#             # --------------------
#             #   导出为 .docx
#             # --------------------
#             doc = Document()
#             doc.add_heading(note.title, level=1)  # 标题
#             doc.add_paragraph(note.content)       # 内容

#             file_stream = io.BytesIO()  # 创建内存流
#             doc.save(file_stream)       # 将生成的 docx 写入内存流
#             file_stream.seek(0)         # 指针回到开头

#             safe_file_name = escape_uri_path(f"{file_name_base}.docx")

#             response = HttpResponse(
#                 file_stream.getvalue(),
#                 content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
#             )
#             response['Content-Disposition'] = f'attachment; filename*=UTF-8\'\'{safe_file_name}'
#             return response

#         else:
#             # --------------------
#             #   默认为 .txt
#             # --------------------
#             file_content = f"{note.title}\n\n{note.content}"
#             safe_file_name = escape_uri_path(f"{file_name_base}.txt")

#             response = HttpResponse(file_content, content_type='text/plain; charset=UTF-8')
#             response['Content-Disposition'] = f'attachment; filename*=UTF-8\'\'{safe_file_name}'
#             return response

#     return JsonResponse({"code": 1, "msg": "无效的请求方法"})
@csrf_exempt
def export_note_view(request):
    """
    导出笔记内容为文件的视图，可选择txt或docx格式
    请求方法: GET
    请求参数:
        note_id (必填)  : 笔记ID
        format (选填)   : 导出格式, 可选 'txt' 或 'docx'，默认为 'txt'
    """
    print("request.method =", request.method)
    if request.method == 'GET':
        # 验证并刷新 token
        try:
            token = verify_and_refresh_token(request)
            access_token = AccessToken(token)
            user_id = access_token['user_id']  # 从 token 中获取 user_id
        except Exception as e:
            return JsonResponse({"code": 1, "msg": f"Token 验证失败: {str(e)}"})

        # 获取前端传递的 note_id 和导出格式
        note_id = request.GET.get('note_id')
        export_format = request.GET.get('format', 'txt').lower()  # 默认 txt

        if not note_id:
            return JsonResponse({"code": 1, "msg": "缺少参数：note_id"})

        # 查询指定笔记
        try:
            note = Note.objects.get(id=note_id, user_id=user_id)
        except Note.DoesNotExist:
            return JsonResponse({"code": 1, "msg": "笔记不存在或无权限导出"})

        # 准备文件名
        file_name_base = note.title.strip() if note.title.strip() else "未命名"
        # 为了避免不合法的文件名，可以再做一次简单处理
        file_name_base = file_name_base.replace('/', '_').replace('\\', '_').replace(':', '-')

        # 根据导出格式进行处理
        if export_format == 'docx':
            # --------------------
            #   导出为 .docx
            # --------------------
            doc = Document()
            doc.add_heading(note.title, level=1)  # 标题
            doc.add_paragraph(note.content)       # 内容

            file_stream = io.BytesIO()  # 创建内存流
            doc.save(file_stream)       # 将生成的 docx 写入内存流
            file_stream.seek(0)         # 指针回到开头

            safe_file_name = escape_uri_path(f"{file_name_base}.docx")

            response = HttpResponse(
                file_stream.getvalue(),
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            response['Content-Disposition'] = f'attachment; filename*=UTF-8\'\'{safe_file_name}'
            return response

        else:
            # --------------------
            #   默认为 .txt
            # --------------------
            file_content = f"{note.title}\n\n{note.content}"
            safe_file_name = escape_uri_path(f"{file_name_base}.txt")

            response = HttpResponse(file_content, content_type='text/plain; charset=UTF-8')
            response['Content-Disposition'] = f'attachment; filename*=UTF-8\'\'{safe_file_name}'
            return response

    return JsonResponse({"code": 1, "msg": "无效的请求方法"})



@csrf_exempt
def import_note_view(request):
    """
    从上传的文件导入笔记内容，支持 txt 或 docx 格式
    请求方法: POST
    请求参数 (multipart/form-data):
        file       (必填): 上传的文件，支持 txt 或 docx
        title      (选填): 如果不填，则尝试用文件名作为标题
        folder_id  (必填): 笔记所属文件夹 ID
        format     (选填): 指定文件格式，如果不提供，则根据文件后缀名推断
    返回:
        - 创建好的笔记信息 或 错误信息
    """

    if request.method == 'POST':
        # 1. 验证并刷新 token
        try:
            token = verify_and_refresh_token(request)
            access_token = AccessToken(token)
            user_id = access_token['user_id']  # 从 token 中获取 user_id
        except Exception as e:
            return JsonResponse({"code": 1, "msg": f"Token 验证失败: {str(e)}"})

        # 2. 获取上传的文件
        uploaded_file = request.FILES.get('file')
        if not uploaded_file:
            return JsonResponse({"code": 1, "msg": "缺少上传文件 (file)"})

        # 3. 获取其他参数
        title = request.POST.get('title', '').strip()
        folder_id = request.POST.get('folder_id')
        if folder_id is None or folder_id == '':
            return JsonResponse({"code": 1, "msg": "缺少文件夹id"})

        # 4. 获取或推断文件格式
        requested_format = request.POST.get('format', '').lower().strip()
        if not requested_format:
            # 如果前端没传 format, 则根据文件后缀名推断
            file_extension = os.path.splitext(uploaded_file.name)[1].lower()
            if file_extension in ['.docx']:
                requested_format = 'docx'
            else:
                requested_format = 'txt'

        # 如果没有传递 title, 默认使用文件名（去掉后缀）
        if not title:
            # 去除后缀作为标题
            title = os.path.splitext(uploaded_file.name)[0]

        # 5. 读取文件并解析内容
        content = ""
        try:
            if requested_format == 'docx':
                # 解析 docx
                file_stream = io.BytesIO(uploaded_file.read())  # 读取文件内容到内存
                doc = Document(file_stream)
                
                # 将 docx 内的所有段落拼接为一个字符串
                paragraphs = [p.text for p in doc.paragraphs]
                content = "\n".join(paragraphs)
            else:
                # 解析 txt
                content = uploaded_file.read().decode('utf-8', errors='ignore')
        except Exception as e:
            return JsonResponse({"code": 1, "msg": f"文件解析失败: {str(e)}"})

        # 6. 在数据库中创建笔记
        note = Note.objects.create(
            user_id=user_id,
            title=title,
            content=content,
            folder_id=folder_id
        )

        # 7. 返回创建好的笔记信息
        return JsonResponse({
            "code": 0,
            "msg": "笔记导入成功",
            "data": {
                "id": note.id,
                "user_id": note.user_id,
                "title": note.title,
                "content": note.content,
                "folder_id": note.folder_id,
                "created_at": note.created_at.strftime("%Y-%m-%d %H:%M:%S"),
                "updated_at": note.updated_at.strftime("%Y-%m-%d %H:%M:%S"),
                "deleted_at": note.deleted_at.strftime("%Y-%m-%d %H:%M:%S") if note.deleted_at else None,
            }
        })

    # 请求方法不是 POST 返回错误
    return JsonResponse({"code": 1, "msg": "无效的请求方法"})